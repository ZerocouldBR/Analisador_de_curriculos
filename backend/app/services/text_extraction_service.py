"""
Servico avancado de extracao de texto de documentos

Suporta:
- PDF (via pdfplumber + OCR avancado para PDFs escaneados)
- DOCX (via python-docx com extracao de imagens embutidas)
- TXT, RTF (leitura direta)
- Imagens (via Tesseract OCR com preprocessamento avancado)
  - JPEG, PNG, BMP, TIFF, GIF
  - WEBP, HEIC/HEIF (com conversao automatica)
- HTML (via BeautifulSoup)

Melhorias de OCR:
- Preprocessamento de imagem (binarizacao, deskew, denoising, contraste)
- Deteccao de paginas mistas (texto + imagem) em PDFs
- Score de confianca do OCR
- Resolucao adaptativa
- Suporte multi-idioma com deteccao automatica
- Pipeline de multi-tentativa com diferentes configuracoes PSM
- Extracao de imagens embutidas em DOCX para OCR
"""
import io
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass, field

# PDF
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# OCR (Tesseract)
try:
    from PIL import Image, ImageFilter, ImageEnhance, ImageOps
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# OpenCV para preprocessamento avancado
try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except (ImportError, AttributeError, Exception):
    CV2_AVAILABLE = False
    np = None  # type: ignore

# HTML
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


@dataclass
class OCRResult:
    """Resultado de OCR com metadados de qualidade"""
    text: str
    confidence: float  # 0.0 a 1.0
    language: str
    pages_processed: int
    pages_with_ocr: int
    pages_with_text: int
    warnings: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class TextExtractionService:
    """
    Servico para extracao de texto de diferentes formatos de documento

    Suporta:
    - PDF (via pdfplumber com OCR avancado)
    - DOCX (via python-docx)
    - TXT, RTF (leitura direta)
    - Imagens (via Tesseract OCR com preprocessamento)
    - HTML (via BeautifulSoup)
    """

    # Valores carregados de settings (sem hardcoding)
    @staticmethod
    def _ocr_resolutions():
        from app.core.config import settings
        return settings.ocr_resolutions

    @staticmethod
    def _min_text_chars():
        from app.core.config import settings
        return settings.ocr_min_text_chars

    @staticmethod
    def _min_ocr_confidence():
        from app.core.config import settings
        return settings.ocr_min_confidence

    @staticmethod
    def _ocr_languages():
        from app.core.config import settings
        return settings.ocr_languages

    @staticmethod
    def _ocr_psm_modes():
        from app.core.config import settings
        return settings.ocr_psm_modes

    @staticmethod
    def _ocr_max_images():
        from app.core.config import settings
        return settings.ocr_max_images_per_docx

    # MIME types de imagem suportados (incluindo formatos modernos)
    SUPPORTED_IMAGE_MIMES = {
        "image/jpeg", "image/png", "image/bmp", "image/tiff",
        "image/gif", "image/webp", "image/heic", "image/heif",
        "image/x-icon", "image/svg+xml",
    }

    # PSM modes para tentativas progressivas de OCR
    PSM_MODES = [6, 3, 4, 1]  # 6=bloco, 3=auto, 4=coluna, 1=auto com OSD

    @staticmethod
    def extract_text(file_path: str, mime_type: str) -> str:
        """
        Extrai texto de um arquivo baseado no tipo MIME

        Args:
            file_path: Caminho para o arquivo
            mime_type: Tipo MIME do arquivo

        Returns:
            Texto extraido

        Raises:
            ValueError: Se formato nao suportado ou biblioteca nao disponivel
        """
        path = Path(file_path)

        if mime_type == "application/pdf":
            return TextExtractionService._extract_from_pdf(path)

        elif mime_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]:
            return TextExtractionService._extract_from_docx(path)

        elif mime_type in [
            "application/vnd.oasis.opendocument.text",  # ODT
            "application/rtf",
        ]:
            return TextExtractionService._extract_from_txt(path)

        elif mime_type == "text/plain":
            return TextExtractionService._extract_from_txt(path)

        elif mime_type in ["text/html", "application/xhtml+xml"]:
            return TextExtractionService._extract_from_html(path)

        elif mime_type.startswith("image/") or mime_type in TextExtractionService.SUPPORTED_IMAGE_MIMES:
            return TextExtractionService._extract_from_image(path)

        else:
            # Tentar detectar pelo conteudo do arquivo como fallback
            text = TextExtractionService._try_fallback_extraction(path)
            if text:
                return text
            raise ValueError(f"Formato nao suportado: {mime_type}")

    @staticmethod
    def extract_text_with_metadata(file_path: str, mime_type: str) -> OCRResult:
        """
        Extrai texto com metadados detalhados de qualidade

        Args:
            file_path: Caminho para o arquivo
            mime_type: Tipo MIME do arquivo

        Returns:
            OCRResult com texto e metadados
        """
        path = Path(file_path)

        if mime_type == "application/pdf":
            return TextExtractionService._extract_from_pdf_advanced(path)

        elif mime_type.startswith("image/"):
            return TextExtractionService._extract_from_image_advanced(path)

        else:
            text = TextExtractionService.extract_text(file_path, mime_type)
            return OCRResult(
                text=text,
                confidence=1.0,
                language=TextExtractionService.detect_language(text),
                pages_processed=1,
                pages_with_ocr=0,
                pages_with_text=1
            )

    # ================================================================
    # PDF Extraction - Enhanced
    # ================================================================

    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        """Extrai texto de PDF com suporte a paginas mistas (texto + imagem)"""
        if not PDF_AVAILABLE:
            raise ValueError("pdfplumber nao esta instalado. Execute: pip install pdfplumber")

        result = TextExtractionService._extract_from_pdf_advanced(file_path)
        return result.text

    @staticmethod
    def _extract_from_pdf_advanced(file_path: Path) -> OCRResult:
        """
        Extracao avancada de PDF com tratamento de paginas mistas

        Para cada pagina:
        1. Tenta extracao direta de texto
        2. Se texto insuficiente, aplica OCR com preprocessamento
        3. Combina resultados de ambos os metodos
        """
        if not PDF_AVAILABLE:
            raise ValueError("pdfplumber nao esta instalado")

        all_text = []
        total_confidence = 0.0
        pages_with_ocr = 0
        pages_with_text = 0
        warnings = []
        page_results = []

        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages):
                    page_text = ""

                    # Tentar extracao direta de texto (com configuracoes otimizadas)
                    try:
                        # Usar configuracao otimizada para extrair texto
                        extract_settings = {
                            "x_tolerance": 3,
                            "y_tolerance": 3,
                            "keep_blank_chars": False,
                        }
                        direct_text = page.extract_text(
                            x_tolerance=extract_settings["x_tolerance"],
                            y_tolerance=extract_settings["y_tolerance"],
                        ) or ""
                        direct_text = direct_text.strip()
                    except Exception as e:
                        direct_text = ""
                        warnings.append(f"Erro texto direto pagina {page_num + 1}: {str(e)}")

                    # Extrair tabelas da pagina (pdfplumber e bom nisso)
                    table_text = ""
                    try:
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row in table:
                                    if row:
                                        cells = [str(cell).strip() for cell in row if cell]
                                        if cells:
                                            table_text += " | ".join(cells) + "\n"
                    except Exception:
                        pass

                    # Decidir se precisa OCR
                    has_sufficient_text = len(direct_text) >= TextExtractionService._min_text_chars()
                    page_confidence = 1.0

                    if has_sufficient_text:
                        page_text = direct_text
                        pages_with_text += 1
                        page_confidence = 1.0
                    elif OCR_AVAILABLE:
                        # Aplicar OCR com preprocessamento na pagina inteira
                        ocr_text, ocr_confidence = TextExtractionService._ocr_pdf_page(
                            page, page_num
                        )

                        if ocr_text and len(ocr_text.strip()) > len(direct_text):
                            page_text = ocr_text
                            page_confidence = ocr_confidence / 100.0
                            pages_with_ocr += 1

                            if ocr_confidence < TextExtractionService._min_ocr_confidence():
                                warnings.append(
                                    f"Pagina {page_num + 1}: OCR com baixa confianca "
                                    f"({ocr_confidence:.1f}%)"
                                )
                        else:
                            page_text = direct_text
                            if not page_text:
                                warnings.append(
                                    f"Pagina {page_num + 1}: nenhum texto extraido"
                                )
                    else:
                        page_text = direct_text

                    # Extrair imagens embutidas na pagina e fazer OCR nelas
                    if OCR_AVAILABLE:
                        try:
                            embedded_img_texts = TextExtractionService._extract_images_from_pdf_page(
                                page, page_num
                            )
                            if embedded_img_texts:
                                page_text += "\n" + "\n".join(embedded_img_texts)
                                if not has_sufficient_text:
                                    pages_with_ocr += 1
                        except Exception as img_err:
                            logger.debug(f"Erro ao extrair imagens da pagina {page_num + 1}: {img_err}")

                    # Combinar texto direto + tabelas
                    combined = page_text
                    if table_text and table_text not in combined:
                        combined += "\n\n[TABELA]\n" + table_text

                    if combined.strip():
                        all_text.append(combined.strip())

                    total_confidence += page_confidence
                    page_results.append({
                        "page": page_num + 1,
                        "method": "text" if has_sufficient_text else "ocr",
                        "confidence": page_confidence,
                        "chars": len(combined)
                    })

            final_text = "\n\n".join(all_text).strip()
            avg_confidence = total_confidence / max(total_pages, 1)

            # Detectar idioma
            language = TextExtractionService.detect_language(final_text)

            return OCRResult(
                text=final_text,
                confidence=avg_confidence,
                language=language,
                pages_processed=total_pages,
                pages_with_ocr=pages_with_ocr,
                pages_with_text=pages_with_text,
                warnings=warnings,
                metadata={
                    "file": str(file_path),
                    "page_results": page_results,
                    "extraction_method": "hybrid" if pages_with_ocr > 0 else "direct"
                }
            )

        except Exception as e:
            raise ValueError(f"Erro ao extrair texto do PDF: {str(e)}")

    @staticmethod
    def _ocr_pdf_page(page, page_num: int) -> Tuple[str, float]:
        """
        Aplica OCR em uma pagina do PDF com preprocessamento

        Tenta multiplas resolucoes e tecnicas de preprocessamento.
        Para cada resolucao, tenta com e sem preprocessamento para
        capturar o melhor resultado possivel.

        Returns:
            Tuple (texto_extraido, confianca_media)
        """
        best_text = ""
        best_confidence = 0.0

        for resolution in TextExtractionService._ocr_resolutions():
            try:
                # Converter pagina para imagem
                page_image = page.to_image(resolution=resolution).original

                # Tentar com preprocessamento
                processed_image = TextExtractionService._preprocess_image(page_image)

                # OCR com dados de confianca
                ocr_data = pytesseract.image_to_data(
                    processed_image,
                    lang=TextExtractionService._ocr_languages(),
                    output_type=pytesseract.Output.DICT,
                    config='--oem 3 --psm 6'
                )

                # Extrair texto e calcular confianca
                text_parts = []
                confidences = []

                for i, conf in enumerate(ocr_data['conf']):
                    word = ocr_data['text'][i].strip()
                    if word and conf != -1:
                        text_parts.append(word)
                        confidences.append(float(conf))

                if text_parts:
                    avg_conf = sum(confidences) / len(confidences) if confidences else 0

                    # Reconstruir com quebras de linha baseado em blocos
                    text = TextExtractionService._reconstruct_text_from_ocr_data(ocr_data)

                    if avg_conf > best_confidence and len(text) > len(best_text):
                        best_text = text
                        best_confidence = avg_conf

                # Tambem tentar sem preprocessamento (imagem original)
                # pois alguns PDFs com texto digital sao prejudicados pelo preprocessamento
                try:
                    ocr_data_raw = pytesseract.image_to_data(
                        page_image,
                        lang=TextExtractionService._ocr_languages(),
                        output_type=pytesseract.Output.DICT,
                        config='--oem 3 --psm 3'
                    )
                    raw_parts = []
                    raw_confs = []
                    for i, conf in enumerate(ocr_data_raw['conf']):
                        word = ocr_data_raw['text'][i].strip()
                        if word and conf != -1:
                            raw_parts.append(word)
                            raw_confs.append(float(conf))

                    if raw_parts:
                        raw_avg = sum(raw_confs) / len(raw_confs)
                        raw_text = TextExtractionService._reconstruct_text_from_ocr_data(ocr_data_raw)
                        raw_quality = len(raw_text.strip()) * (raw_avg / 100.0)
                        best_quality = len(best_text.strip()) * (best_confidence / 100.0)
                        if raw_quality > best_quality:
                            best_text = raw_text
                            best_confidence = raw_avg
                except Exception:
                    pass

                # Se confianca boa, nao precisa tentar outras resolucoes
                from app.core.config import settings as _settings
                if best_confidence >= _settings.ocr_good_confidence_threshold:
                    break

            except Exception as e:
                logger.warning(f"OCR falhou na pagina {page_num + 1} com {resolution}dpi: {e}")
                continue

        return best_text, best_confidence

    @staticmethod
    def _extract_images_from_pdf_page(page, page_num: int) -> list:
        """
        Extrai imagens embutidas em uma pagina do PDF e aplica OCR nelas.
        Isso captura logos, fotos e imagens com texto que nao sao capturadas
        pela extracao de texto normal.
        """
        image_texts = []

        try:
            images = page.images
            if not images:
                return image_texts

            for img_idx, img in enumerate(images[:10]):  # Limitar a 10 imagens por pagina
                try:
                    # Extrair regiao da imagem da pagina
                    x0 = img.get("x0", 0)
                    y0 = img.get("top", 0)
                    x1 = img.get("x1", page.width)
                    y1 = img.get("bottom", page.height)

                    # Ignorar imagens muito pequenas (icones, decoracao)
                    img_width = x1 - x0
                    img_height = y1 - y0
                    if img_width < 50 or img_height < 30:
                        continue

                    # Converter a regiao da pagina para imagem
                    cropped = page.crop((x0, y0, x1, y1))
                    page_image = cropped.to_image(resolution=300).original

                    # Preprocessar e OCR
                    processed = TextExtractionService._preprocess_image(page_image)
                    ocr_text = pytesseract.image_to_string(
                        processed,
                        lang=TextExtractionService._ocr_languages(),
                        config='--oem 3 --psm 6'
                    ).strip()

                    if len(ocr_text) > 20:  # Ignorar textos muito curtos
                        image_texts.append(ocr_text)

                except Exception as e:
                    logger.debug(f"Erro OCR imagem {img_idx} pagina {page_num + 1}: {e}")
                    continue

        except Exception as e:
            logger.debug(f"Erro ao listar imagens pagina {page_num + 1}: {e}")

        return image_texts

    @staticmethod
    def _reconstruct_text_from_ocr_data(ocr_data: dict) -> str:
        """Reconstroi texto preservando estrutura de paragrafos e linhas"""
        lines = {}
        for i, word in enumerate(ocr_data['text']):
            word = word.strip()
            if not word or ocr_data['conf'][i] == -1:
                continue

            block_num = ocr_data['block_num'][i]
            line_num = ocr_data['line_num'][i]
            key = (block_num, line_num)

            if key not in lines:
                lines[key] = []
            lines[key].append(word)

        # Construir texto com estrutura
        result = []
        prev_block = None

        for (block_num, line_num) in sorted(lines.keys()):
            if prev_block is not None and block_num != prev_block:
                result.append("")  # Linha vazia entre blocos
            result.append(" ".join(lines[(block_num, line_num)]))
            prev_block = block_num

        return "\n".join(result)

    # ================================================================
    # Image Preprocessing for OCR
    # ================================================================

    @staticmethod
    def _preprocess_image(image: Image.Image) -> Image.Image:
        """
        Aplica pipeline de preprocessamento para melhorar OCR

        Pipeline:
        1. Converter para escala de cinza
        2. Ajustar contraste
        3. Remover ruido (denoise)
        4. Binarizacao adaptativa
        5. Correcao de inclinacao (deskew)
        """
        if CV2_AVAILABLE:
            return TextExtractionService._preprocess_with_opencv(image)
        else:
            return TextExtractionService._preprocess_with_pillow(image)

    @staticmethod
    def _preprocess_with_opencv(image: Image.Image) -> Image.Image:
        """Preprocessamento avancado usando OpenCV"""
        # Converter PIL para numpy array
        img_array = np.array(image)

        # Converter para escala de cinza se necessario
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        # Redimensionar se imagem muito pequena
        height, width = gray.shape
        if height < 500 or width < 500:
            scale = max(1000 / width, 1000 / height)
            gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

        # Remover ruido com filtro bilateral (preserva bordas)
        denoised = cv2.bilateralFilter(gray, 9, 75, 75)

        # Ajustar contraste com CLAHE (Contrast Limited Adaptive Histogram Equalization)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        contrast = clahe.apply(denoised)

        # Binarizacao adaptativa (Gaussian)
        binary = cv2.adaptiveThreshold(
            contrast, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            blockSize=15,
            C=8
        )

        # Correcao de inclinacao (deskew)
        binary = TextExtractionService._deskew_opencv(binary)

        # Operacoes morfologicas para limpar ruido residual
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
        binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

        # Remover bordas escuras / artefatos de scan
        binary = TextExtractionService._remove_scan_borders(binary)

        # Converter de volta para PIL
        return Image.fromarray(binary)

    @staticmethod
    def _deskew_opencv(image: "np.ndarray") -> "np.ndarray":
        """Corrige inclinacao da imagem usando OpenCV"""
        try:
            # Detectar angulo de inclinacao
            coords = np.column_stack(np.where(image < 128))
            if len(coords) < 100:
                return image

            angle = cv2.minAreaRect(coords)[-1]

            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle

            # So corrigir se inclinacao significativa mas nao extrema
            if abs(angle) > 0.5 and abs(angle) < 15:
                (h, w) = image.shape[:2]
                center = (w // 2, h // 2)
                rotation_matrix = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(
                    image, rotation_matrix, (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE
                )
                return rotated

        except Exception as e:
            logger.debug(f"Deskew falhou: {e}")

        return image

    @staticmethod
    def _remove_scan_borders(image: "np.ndarray") -> "np.ndarray":
        """Remove bordas escuras tipicas de documentos escaneados"""
        try:
            h, w = image.shape[:2]
            border_size = min(h, w) // 50  # ~2% da dimensao menor

            if border_size < 3:
                return image

            # Verificar se bordas sao predominantemente escuras
            top = image[:border_size, :].mean()
            bottom = image[-border_size:, :].mean()
            left = image[:, :border_size].mean()
            right = image[:, -border_size:].mean()

            # Se borda e escura (< 128 media), preencher com branco
            if top < 128:
                image[:border_size, :] = 255
            if bottom < 128:
                image[-border_size:, :] = 255
            if left < 128:
                image[:, :border_size] = 255
            if right < 128:
                image[:, -border_size:] = 255

        except Exception:
            pass

        return image

    @staticmethod
    def _preprocess_with_pillow(image: Image.Image) -> Image.Image:
        """Preprocessamento usando apenas Pillow (fallback sem OpenCV)"""
        # Converter para escala de cinza
        if image.mode != 'L':
            image = image.convert('L')

        # Aumentar contraste
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)

        # Aumentar nitidez
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(2.0)

        # Remover ruido com filtro mediano
        image = image.filter(ImageFilter.MedianFilter(size=3))

        # Binarizacao com threshold de Otsu simulado
        # Calcular threshold otimo
        histogram = image.histogram()
        total_pixels = sum(histogram)
        threshold = TextExtractionService._otsu_threshold(histogram, total_pixels)

        # Aplicar threshold
        image = image.point(lambda x: 255 if x > threshold else 0, '1')

        return image

    @staticmethod
    def _otsu_threshold(histogram: list, total: int) -> int:
        """Calcula threshold otimo usando metodo de Otsu"""
        sum_total = sum(i * histogram[i] for i in range(256))
        sum_bg = 0
        weight_bg = 0
        max_variance = 0
        threshold = 0

        for t in range(256):
            weight_bg += histogram[t]
            if weight_bg == 0:
                continue

            weight_fg = total - weight_bg
            if weight_fg == 0:
                break

            sum_bg += t * histogram[t]
            mean_bg = sum_bg / weight_bg
            mean_fg = (sum_total - sum_bg) / weight_fg

            variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2

            if variance > max_variance:
                max_variance = variance
                threshold = t

        return threshold

    # ================================================================
    # Image Extraction - Enhanced
    # ================================================================

    @staticmethod
    def _extract_from_image(file_path: Path) -> str:
        """Extrai texto de imagem usando OCR com preprocessamento"""
        result = TextExtractionService._extract_from_image_advanced(file_path)
        return result.text

    @staticmethod
    def _convert_image_to_pil(file_path: Path) -> Image.Image:
        """
        Converte imagem de qualquer formato suportado para PIL Image.
        Suporta WEBP, HEIC, HEIF alem dos formatos tradicionais.
        """
        suffix = file_path.suffix.lower()

        try:
            if suffix in ('.heic', '.heif'):
                # Tentar usar pillow-heif para HEIC/HEIF
                try:
                    import pillow_heif
                    heif_file = pillow_heif.read_heif(str(file_path))
                    image = Image.frombytes(
                        heif_file.mode, heif_file.size, heif_file.data,
                        "raw", heif_file.mode, heif_file.stride,
                    )
                    return image
                except ImportError:
                    logger.warning("pillow-heif nao disponivel. Tentando conversao com OpenCV.")
                    if CV2_AVAILABLE:
                        img_array = cv2.imread(str(file_path))
                        if img_array is not None:
                            img_rgb = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
                            return Image.fromarray(img_rgb)
                    raise ValueError(
                        "Formato HEIC/HEIF requer pillow-heif. "
                        "Execute: pip install pillow-heif"
                    )
            else:
                image = Image.open(file_path)
                # Converter WEBP e outros formatos com canal alpha para RGB
                if image.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', image.size, (255, 255, 255))
                    if image.mode == 'P':
                        image = image.convert('RGBA')
                    background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                    image = background
                elif image.mode != 'RGB' and image.mode != 'L':
                    image = image.convert('RGB')
                return image

        except Exception as e:
            raise ValueError(f"Erro ao abrir imagem {file_path.name}: {str(e)}")

    @staticmethod
    def _extract_from_image_advanced(file_path: Path) -> OCRResult:
        """
        Extracao avancada de imagem com metadados.
        Suporta multiplos formatos e usa multi-tentativa com diferentes PSM modes.
        """
        if not OCR_AVAILABLE:
            raise ValueError(
                "Bibliotecas de OCR nao disponiveis. "
                "Execute: pip install pytesseract pillow\n"
                "E instale o Tesseract: https://github.com/tesseract-ocr/tesseract"
            )

        try:
            image = TextExtractionService._convert_image_to_pil(file_path)
            original_size = image.size

            # Preprocessar imagem
            processed = TextExtractionService._preprocess_image(image)

            # Multi-tentativa com diferentes PSM modes para melhor resultado
            best_text = ""
            best_confidence = 0.0
            best_psm = 6

            for psm in TextExtractionService._ocr_psm_modes():
                try:
                    ocr_data = pytesseract.image_to_data(
                        processed,
                        lang=TextExtractionService._ocr_languages(),
                        output_type=pytesseract.Output.DICT,
                        config=f'--oem 3 --psm {psm}'
                    )

                    text = TextExtractionService._reconstruct_text_from_ocr_data(ocr_data)

                    # Calcular confianca com indice correto
                    confidences = []
                    for i, conf in enumerate(ocr_data['conf']):
                        if conf != -1 and ocr_data['text'][i].strip():
                            confidences.append(float(conf))

                    avg_conf = sum(confidences) / len(confidences) if confidences else 0

                    # Preferir resultado com mais texto e melhor confianca
                    text_quality = len(text.strip()) * (avg_conf / 100.0)
                    best_quality = len(best_text.strip()) * (best_confidence / 100.0)

                    if text_quality > best_quality:
                        best_text = text
                        best_confidence = avg_conf
                        best_psm = psm

                    # Se confianca boa, nao precisa tentar outros PSMs
                    if avg_conf >= 75 and len(text.strip()) > 50:
                        break

                except Exception as e:
                    logger.debug(f"OCR PSM {psm} falhou: {e}")
                    continue

            language = TextExtractionService.detect_language(best_text)

            warnings = []
            if best_confidence < TextExtractionService._min_ocr_confidence():
                warnings.append(f"Baixa confianca do OCR: {best_confidence:.1f}%")
            if not best_text.strip():
                warnings.append("Nenhum texto detectado na imagem")

            return OCRResult(
                text=best_text.strip(),
                confidence=best_confidence / 100.0,
                language=language,
                pages_processed=1,
                pages_with_ocr=1,
                pages_with_text=0,
                warnings=warnings,
                metadata={
                    "file": str(file_path),
                    "image_size": original_size,
                    "ocr_engine": "tesseract",
                    "psm_used": best_psm,
                    "format": file_path.suffix.lower(),
                }
            )

        except Exception as e:
            raise ValueError(f"Erro ao extrair texto da imagem: {str(e)}")

    # ================================================================
    # DOCX, TXT, HTML Extraction
    # ================================================================

    @staticmethod
    def _extract_from_docx(file_path: Path) -> str:
        """Extrai texto de DOCX incluindo headers, footers, tabelas e imagens embutidas via OCR"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx nao esta instalado. Execute: pip install python-docx")

        try:
            doc = DocxDocument(file_path)

            text = []

            # Extrair paragrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extrair tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))

            # Extrair texto de imagens embutidas via OCR
            if OCR_AVAILABLE:
                image_texts = TextExtractionService._extract_images_from_docx(file_path)
                if image_texts:
                    text.append("\n[TEXTO EXTRAIDO DE IMAGENS EMBUTIDAS]")
                    text.extend(image_texts)

            return "\n".join(text)

        except Exception as e:
            raise ValueError(f"Erro ao extrair texto do DOCX: {str(e)}")

    @staticmethod
    def _extract_images_from_docx(file_path: Path) -> list:
        """Extrai e faz OCR em imagens embutidas dentro de DOCX"""
        import zipfile
        image_texts = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                image_files = [
                    f for f in zip_ref.namelist()
                    if f.startswith('word/media/') and any(
                        f.lower().endswith(ext) for ext in
                        ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif')
                    )
                ]

                for img_file in image_files[:TextExtractionService._ocr_max_images()]:
                    try:
                        img_data = zip_ref.read(img_file)
                        image = Image.open(io.BytesIO(img_data))

                        # So processar imagens grandes o suficiente para conter texto
                        if image.size[0] < 100 or image.size[1] < 50:
                            continue

                        processed = TextExtractionService._preprocess_image(image)
                        ocr_text = pytesseract.image_to_string(
                            processed, lang=TextExtractionService._ocr_languages(),
                            config='--oem 3 --psm 6'
                        ).strip()

                        if len(ocr_text) > 20:  # Ignorar textos muito curtos
                            image_texts.append(ocr_text)

                    except Exception as e:
                        logger.debug(f"Erro OCR em imagem embutida {img_file}: {e}")
                        continue

        except Exception as e:
            logger.debug(f"Erro ao extrair imagens do DOCX: {e}")

        return image_texts

    @staticmethod
    def _extract_from_txt(file_path: Path) -> str:
        """Extrai texto de TXT"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Nao foi possivel detectar encoding do arquivo")

        except Exception as e:
            raise ValueError(f"Erro ao ler arquivo TXT: {str(e)}")

    @staticmethod
    def _extract_from_html(file_path: Path) -> str:
        """Extrai texto de HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text()

            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text

        except Exception as e:
            raise ValueError(f"Erro ao extrair texto do HTML: {str(e)}")

    # ================================================================
    # Text Normalization & Language Detection
    # ================================================================

    @staticmethod
    def normalize_text(text: str) -> str:
        """
        Normaliza texto extraido

        - Remove espacos extras
        - Normaliza quebras de linha
        - Remove caracteres especiais problematicos
        - Corrige artefatos comuns de OCR
        """
        # Remover caracteres de controle (exceto \n e \t)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # Normalizar quebras de linha
        text = re.sub(r'\r\n|\r', '\n', text)

        # Corrigir artefatos comuns de OCR
        text = TextExtractionService._fix_ocr_artifacts(text)

        # Remover multiplas quebras de linha
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remover espacos no inicio/fim de linhas
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        # Remover multiplos espacos
        text = re.sub(r' {2,}', ' ', text)

        return text.strip()

    @staticmethod
    def _fix_ocr_artifacts(text: str) -> str:
        """Corrige artefatos comuns de OCR em texto portugues"""
        # Correcoes comuns de OCR
        corrections = {
            # Caracteres confundidos
            r'\bl\b(?=\s+[a-z])': 'I',  # l isolado -> I
            r'(?<=[a-z])0(?=[a-z])': 'o',  # 0 entre letras -> o
            r'(?<=[a-z])1(?=[a-z])': 'l',  # 1 entre letras -> l
            # Espacos errados em palavras comuns
            r'expe riência': 'experiência',
            r'expe riencia': 'experiencia',
            r'forma ção': 'formação',
            r'forma cao': 'formacao',
            r'educa ção': 'educação',
            r'habili dades': 'habilidades',
            r'certi ficações': 'certificações',
            r'produ ção': 'produção',
            r'logís tica': 'logística',
            r'manu tenção': 'manutenção',
            # Simbolos confundidos
            r'@\s+': '@',  # espaco apos @
            r'\s+@': '@',  # espaco antes de @
        }

        for pattern, replacement in corrections.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

        return text

    @staticmethod
    def _try_fallback_extraction(file_path: Path) -> Optional[str]:
        """
        Tenta extrair texto por deteccao de conteudo quando mime type eh desconhecido.
        Usa magic bytes para detectar o formato real do arquivo.
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(16)

            # PDF magic bytes
            if header[:5] == b'%PDF-':
                if PDF_AVAILABLE:
                    return TextExtractionService._extract_from_pdf(file_path)

            # DOCX/ZIP magic bytes
            if header[:4] == b'PK\x03\x04':
                if DOCX_AVAILABLE:
                    try:
                        return TextExtractionService._extract_from_docx(file_path)
                    except Exception:
                        pass

            # JPEG magic bytes
            if header[:3] == b'\xff\xd8\xff':
                if OCR_AVAILABLE:
                    return TextExtractionService._extract_from_image(file_path)

            # PNG magic bytes
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                if OCR_AVAILABLE:
                    return TextExtractionService._extract_from_image(file_path)

            # Tentar como texto puro
            try:
                return TextExtractionService._extract_from_txt(file_path)
            except Exception:
                pass

        except Exception as e:
            logger.debug(f"Fallback extraction falhou: {e}")

        return None

    @staticmethod
    def detect_language(text: str) -> str:
        """
        Detecta idioma do texto (simplificado)

        Args:
            text: Texto para analise

        Returns:
            Codigo do idioma ('pt', 'en', 'es', etc.)
        """
        # Palavras comuns por idioma
        language_words = {
            'pt': ['de', 'e', 'o', 'a', 'para', 'com', 'em', 'é', 'são', 'como',
                    'experiência', 'formação', 'habilidades', 'empresa', 'trabalho',
                    'produção', 'logística', 'operador', 'auxiliar'],
            'en': ['the', 'and', 'of', 'to', 'for', 'with', 'in', 'is', 'are', 'as',
                    'experience', 'education', 'skills', 'company', 'work',
                    'production', 'logistics', 'operator', 'assistant'],
            'es': ['el', 'la', 'de', 'en', 'y', 'que', 'los', 'del', 'las', 'por',
                    'experiencia', 'formación', 'habilidades', 'empresa', 'trabajo'],
        }

        text_lower = text.lower()
        scores = {}

        for lang, words in language_words.items():
            score = sum(1 for word in words if f' {word} ' in f' {text_lower} ')
            scores[lang] = score

        if not scores or max(scores.values()) == 0:
            return 'unknown'

        return max(scores, key=scores.get)
